import tkinter as tk
from tkinter import *
from tkinter import ttk ,FLAT
from PIL import Image, ImageTk, ImageEnhance
import cv2
import numpy as np
import threading
import os
import time
import docx
from PIL import Image
import ctypes, sys
from utils.utils import *
from utils.cosine_summary import *
## docx2pdf doesn't seem to be working for Ubuntu
# from docx2pdf import convert

import pyaudio
import wave
import multiprocessing

frams = []
sound  = True
CHUNK = 1024
FORMAT = pyaudio.paInt16
CHANNELS = 2
RATE = 44100
RECORD_SECONDS = 1000
WAVE_OUTPUT_FILENAME = "output.wav"


# Since ImageGrab exists with PIL only in MACOS and Windows
try:
    from PIL import ImageGrab
except:
    import pyscreenshot as ImageGrab
    
p = pyaudio.PyAudio()

stream = p.open(format=FORMAT,
                channels=CHANNELS,
                rate=RATE,
                input=True,
                frames_per_buffer=CHUNK)


doc = docx.Document()
summary_doc = docx.Document()
try:
    if not os.path.exists('data'):
        os.makedirs('data')

    # if not created then raise error
except OSError:
    print('Error: Creating directory of data')


VIDEO_SIZE = (1920,1080)
#f = ImageGrab.grab() ### not necessary
#a, b = f.size        ### not necessary
filename="test1.avi"
fourcc = cv2.VideoWriter_fourcc(*'DIVX')
frame_rate = 10
out = cv2.VideoWriter(filename, fourcc, frame_rate, VIDEO_SIZE) # use VIDEO_SIZE
images = []
frames=0

root = tk.Tk()
root.resizable(0, 0)
root.title('Screen Recorder')
root.geometry('+260+70')
#x1 = y1 = x2 = y2 = 0  # not necessary
def show_image(image):
    win = tk.Toplevel()
    win.image = ImageTk.PhotoImage(image)
    tk.Label(win, image=win.image).pack()
    win.grab_set()
    win.wait_window(win)

def area_sel():
    x1 = y1 = x2 = y2 = 0
    roi_image = None

    def on_mouse_down(event):
        nonlocal x1, y1
        x1, y1 = event.x, event.y
        canvas.create_rectangle(x1, y1, x1, y1, outline='red', tag='roi')

    def on_mouse_move(event):
        nonlocal roi_image, x2, y2
        x2, y2 = event.x, event.y
        canvas.delete('roi-image')
        roi_image = image.crop((x1, y1, x2, y2))
        canvas.image = ImageTk.PhotoImage(roi_image)
        canvas.create_image(x1, y1, image=canvas.image, tag=('roi-image'), anchor='nw')
        canvas.coords('roi', x1, y1, x2, y2)
        canvas.lift('roi')

    root.withdraw()
    image = ImageGrab.grab()
    bgimage = ImageEnhance.Brightness(image).enhance(0.5)
    win = tk.Toplevel()
    win.attributes('-fullscreen', 1)
    win.attributes('-topmost', 1)
    canvas = tk.Canvas(win, highlightthickness=0)
    canvas.pack(fill='both', expand=1)
    tkimage = ImageTk.PhotoImage(bgimage)
    canvas.create_image(0, 0, image=tkimage, anchor='nw', tag='images')
    win.bind('<ButtonPress-1>', on_mouse_down)
    win.bind('<B1-Motion>', on_mouse_move)
    win.bind('<ButtonRelease-1>', lambda e: win.destroy())
    win.focus_force()
    win.grab_set()
    win.wait_window(win)
    root.deiconify()  

    if roi_image:
        # p1 = multiprocessing.Process(target=start_recording,args=[(x1, y1, x2, y2)])
        # p2 = multiprocessing.Process(target=audio,args=[])

        # if __name__ == "__main__":
        #     p1
        #     p2.start()
        start_recording((x1, y1, x2, y2)) #calling main function to record screen



def recording_screen(x1, y1, x2, y2):
    global recording
    recording = True

    global frames

    step = 3
    frames_count = 100000
    

    currentframe = 0
    frame_per_second = 10
    frames_captured = 0

    while recording:
        img = ImageGrab.grab(bbox=(x1, y1, x2, y2))
        #frame = np.array(img) # not necessary
        sc = np.array(img)
        sc = cv2.resize(sc, VIDEO_SIZE)
        tkimage.paste(Image.fromarray(sc))
        frame = cv2.cvtColor(sc, cv2.COLOR_RGB2BGR)
        out.write(frame)
        
        
        #------------------------------------------------------------------------------
        if currentframe > (step*frame_per_second):  
            currentframe = 0

            images.append(frame)
            frames+=1
            # print('Image_data shape:', np.array(images).shape)

            if frames_captured>=1:

                #  Difference in images part
                i1 = Image.fromarray(images[frames_captured-1])
                i2 = Image.fromarray(images[frames_captured])
                assert i1.mode == i2.mode, "Different kinds of images."
                assert i1.size == i2.size, "Different sizes."
                
                pairs = zip(i1.getdata(), i2.getdata())
                if len(i1.getbands()) == 1:
                    # for gray-scale jpegs
                    dif = sum(abs(p1-p2) for p1,p2 in pairs)
                else:
                    dif = sum(abs(c1-c2) for p1,p2 in pairs for c1,c2 in zip(p1,p2))
                
                ncomponents = i1.size[0] * i1.size[1] * 3
                print ("Difference (percentage):", (dif / 255.0 * 100) / ncomponents)
                percentage = (dif / 255.0 * 100) / ncomponents
                if percentage > 0.7:
                    name = './data/frame' + str(frames_captured-1) + '.jpg'

                    #Saving the image
                    cv2.imwrite(name, images[frames_captured-1])

                    # Adding image to a doc
                    doc.add_picture('./data/frame' + str(frames_captured-1) + '.jpg',width=docx.shared.Inches(6), height = docx.shared.Inches(3))
                    
                    #Removing the image from storage
                    os.remove('./data/frame' + str(frames_captured-1) + '.jpg')


            frames_captured+=1 

            if frames_captured>frames_count-1:
                cv2.imwrite("Last.JPG", images[frames-1])
                doc.add_picture('Last.JPG',width=docx.shared.Inches(6), height = docx.shared.Inches(3))
                os.remove('Last.JPG')
        currentframe += 1   
            
        #------------------------------------------------------------------------------------

def start_recording(region):
    ''' not necessary
    if not out.isOpened():
        out.open(filename, fourcc, frame_rate, VIDEO_SIZE)
    '''
    if __name__ == "__main__":
        p1 = threading.Thread(target=recording_screen, args=region, daemon=True)
        p2 = threading.Thread(target=audio)
        p1.start()
        p2.start()

def stop_recording():
    global recording
    recording = False
    global sound,stream, p
    sound = False

    print("* done recording")
    stream.stop_stream()
    stream.close()
    p.terminate()

    wf = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
    wf.setnchannels(CHANNELS)
    wf.setsampwidth(p.get_sample_size(FORMAT))
    wf.setframerate(RATE)
    wf.writeframes(b''.join(frams))
    wf.close()


    out.release()
    cv2.imwrite("Last.JPG", images[frames-1])
    doc.add_picture('Last.JPG',width=docx.shared.Inches(6), height = docx.shared.Inches(3))
    os.remove('Last.JPG')
    doc.save(name.get() + '.docx')
    endmsg = Label(root, text="File Saved as "+ name.get()+".docx")
    endmsg.pack()
    text_from_speech = text_det(WAVE_OUTPUT_FILENAME)
    summary = cosine_summary()
    summarised_text, ranked_sentences = summary.summariser(text_from_speech)
    print(summarised_text)


def audio() :
    global sound, RATE, CHUNK, RECORD_SECONDS, stream

    print("* recording")

    for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
        if sound :
            data = stream.read(CHUNK)
            frams.append(data)


# def docx_to_pdf():
#     convert(name.get()+".docx")
#     endmsg = Label(root, text="Completed!")
#     endmsg.pack() 


tkimage = ImageTk.PhotoImage(Image.new('RGB', VIDEO_SIZE, (0,0,0)))
name = Entry(root)
name.pack()
name.insert(0, "File Name")
mylabel = Label(root, text="By Team Crewmates")
# w, h = VIDEO_SIZE
# vbox = tk.Label(root, image=tkimage, width=w, height=h, bg='black')
# vbox.pack(pady=10,padx=25)

frame = tk.Frame(root)
frame.pack()

sel_area = ttk.Button(frame, text='Start Recording', width=15, command=area_sel)
sel_area.grid(row=0, column=0)

stp_rec = ttk.Button(frame, text='Stop Recording', width=15, command=stop_recording)
stp_rec.grid(row=0, column=1)

# docx_pdf = ttk.Button(frame, text='Covert to PDF', width=15, command=docx_to_pdf)
# docx_pdf.grid(row=1, column=0)

mylabel.pack()

root.mainloop()